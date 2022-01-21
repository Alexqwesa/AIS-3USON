#!/usr/bin/python3
# -------------------------------------------------------------------------------
# -*- coding: utf-8 -*-
# Name:        AIS 3USON fill templates
# Purpose:
#
# Author:      Savin Alexander Viktorovich aka alexqwesa
# Created:     2019
# Copyright:   Savin Alexander Viktorovich aka alexqwesa
# Licence:     LGPL 3
# This software is licensed under the "LGPLv3" License as described in the "LICENSE" file,
# which should be included with this package. The terms are also available at
# http://www.gnu.org/licenses/lgpl-3.0.html
# -------------------------------------------------------------------------------
import subprocess
from _pydecimal import setcontext, BasicContext, Decimal
from operator import itemgetter
# import win32com.client
# import decimal
from decimal import *
import tempfile
#
# from qtpy.QtCore import QDate, QObject
# from qtpy.QtSql import QSqlTableModel, QSqlQuery
#
import typing

from logic.data_worker import *
# from safe_shared_data import SD
# from ts_models import SQL_DATE_FORMAT, tsQsfpModel
# try:
from thirdparty.num2t4ru import num2text
from docx.oxml.ns import qn

from docx.oxml import OxmlElement
from docx import Document
from docx.text.run import Run
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# except ImportError:
#     critical("ImportError")

# from docx.shared import Inches
from models.ts_models_plus import tsQsfpModel


def dlconv(dl, name, value):
    dl["___" + name] = value
    try:
        dl[name] = "{0:0.2f}".format(value).replace(".", ",")
    except TypeError:
        error("empty value %s", name)
        dl[name] = ""
    return dl[name]


class base_pd(QObject):
    def __init__(self, parent=None, template=None):
        """Constructor for prepare_document"""
        self.doc = None
        self.taglist = {}
        self.table = None
        #############################
        # real init
        # ---------------------------
        super().__init__(parent)
        self.reinit(template)

    def reinit(self, template=None):
        self.doc = None
        if template:
            self.doc = Document(template)
        self.taglist = {}
        self.table = None

    @staticmethod
    def dcr(txt):
        return "$__" + str(txt) + "__"

    @staticmethod
    def cropFIO(s: str, to_left: bool = False):
        # return ""
        if "." in s:
            return s

        def shorter(w: str):
            if "-" in w:
                return w[0].upper() + "-" + w.split("-")[1][0] + "."
            else:
                return w[0].upper() + "."

        #############################
        # normalise
        # ---------------------------
        s = s.lower()
        s = s.strip().replace(" ", chr(30))
        s = s.replace(" -", "-")
        s = s.replace("- ", "-")
        s = s.replace(" '", "-")
        s = s.replace("' ", "-")
        sv = s.split()
        # check vector length
        if len(sv) <= 1:
            return s
        l_name = ""
        f_name = ""
        m_name = ""
        #############################
        # drop unused parts
        # ---------------------------
        if sv[0] in ["де", "дел", "дос", "cент", "ван", "фон", "цу"]:
            sv.remove(sv[0])
        if sv[-1] in ["паша", "хан", "шах", "шейх"]:
            sv.pop()
        #############################
        # m_name
        # ---------------------------
        if sv[-1] in ["оглы", "кызы", "заде"]:
            sv.pop()
            m_name = shorter(sv.pop()[0])
        elif sv[-1][-3:] in ["вич", "вна"]:
            if len(sv) >= 3:
                m_name = shorter(sv.pop())
            else:
                f_name = shorter(sv.pop())
        elif sv[-1][-5:] in ["-оглы", "-кызы", "-заде", "-угли", "-уулы"] or sv[-1][:-4] == "-оол":
            if len(sv) >= 3:
                m_name = shorter(sv.pop())
        elif len(sv) >= 4:
            if sv[-2] in ["ибн", "бен", "бин"]:
                m_name = shorter(sv.pop())
                sv.pop()
        elif len(sv) == 3:
            m_name = shorter(sv.pop())
        #############################
        # l_name
        # ---------------------------
        l_name = sv[0]
        if not f_name:
            f_name = shorter(sv[1])
            # f_name = shorter(sv.pop())
        #############################
        # return
        # ---------------------------
        l_name = l_name[0].upper() + l_name[1:]
        if to_left:
            return f_name + m_name + " " + l_name
        else:
            return l_name + " " + f_name + m_name

    @staticmethod
    def num_to_string(value, retformat=1):
        integral, _, exp = str(PrepareDocument.cut_num(value)).replace(".", ",").partition(',')
        if not exp:
            exp = 0
        if retformat == 1:
            return '{} руб. {} коп.'.format(
                num2text(int(integral)), int(exp))
        elif retformat == 2:
            return '{} рублей {} копеек'.format(
                num2text(int(integral)), int(exp))
        elif retformat == 3:
            return num2text(int(integral)), int(exp)
        # num2text(int(exp)))

    @staticmethod
    def cut_num(value):
        # getcontext().prec = 2
        setcontext(BasicContext)
        try:
            value = Decimal(value)
            # q = Decimal(10) ** -2
            return Decimal(value).quantize(Decimal('0.01'))
        except:
            return value


# @auto_reload_class_code(try_wrapper)
# @auto_reload_class_code(reloader)
class PrepareDocument(base_pd):
    """prepare documents from templates"""
    # TODO: add global lock
    __instance: QObject = None

    def __new__(cls, parent=None, template=None):
        inst = PrepareDocument.__instance
        if inst is None:
            PrepareDocument.__instance = QObject.__new__(cls, parent=None, template=None)
            inst = PrepareDocument.__instance
        inst.reinit(parent, template)
        return inst

    @staticmethod
    def console_connect():
        #############################
        # setup dbconnection
        # ---------------------------
        # SD.read_config()
        tr = PrepareDocument.tr
        if not SD.user:
            SD.set_user(input(tr("Введите имя пользователя:")))
        if not SD.password:
            SD.set_password(input(tr("Введите пароль пользователя:")))
        ret = SD.get_db
        return ret

    @staticmethod
    def fill_tags(rep_list, document, search_tables=False):
        """
        in document replace fromtxt to totxt
        """
        for fromtxt, totxt in rep_list.items():
            underline = False
            strike = False
            if isinstance(totxt, list):
                to_txt = totxt[0]
                underline = totxt[1]
                strike = totxt[2]
            else:
                to_txt = totxt
            # try:
            run = PrepareDocument.replace_in_doc(document, str(fromtxt), str(to_txt), search_tables)
            if underline:
                run.underline = underline
            if strike:
                run.font.strike = strike
            # except Exception as uex:
            #     error(uex.args[0])

    @staticmethod
    def replace_in_doc(document, fromtxt, totxt, search_tables=False) -> Run:
        """
        find and coalesce fromtxt in document, replace it to totxt
        and return Run object with totxt
        """
        ret = None
        for par in document.paragraphs:
            #############################
            # go through paragraphs
            # ---------------------------
            ret = PrepareDocument.replace_in_par(par, fromtxt, totxt)
            # ret = False
            if ret:
                return ret
        if search_tables:
            for table in document.tables:
                for cell in table._cells:
                    for par in cell.paragraphs:
                        ret = PrepareDocument.replace_in_par(par, fromtxt, totxt)
                        if ret:
                            return ret
        # raise Exception("Not replaced %s" % fromtxt)

    @staticmethod
    def replace_in_par(par, fromtxt, totxt) -> typing.Union[Run, None]:
        #############################
        # go through runs of paragraph
        # ---------------------------
        ltag = len(fromtxt)
        for i, run in enumerate(par.runs):
            # print(run.text)
            #############################
            # found full match
            # ---------------------------
            # print(run.text)
            lr = len(par.runs)
            if fromtxt in run.text:
                run.text = run.text.replace(fromtxt, totxt)
                return run
                # print (run.text)
            #############################
            # found partial match
            # ---------------------------
            elif (run.text in fromtxt[0:len(run.text)]
                  and lr > i + 1):
                if len(par.runs[i + 1].text) + len(run.text) == ltag:
                    if run.text + par.runs[i + 1].text in fromtxt:
                        run.text = totxt
                        par.runs[i + 1].text = ""
                        return run
                        # print(run.text)
                elif len(par.runs[i + 1].text) + len(run.text) > ltag:
                    rrlen = ltag - len(run.text)
                    if run.text + par.runs[i + 1].text[0:rrlen] in fromtxt:
                        run.text = totxt
                        par.runs[i + 1].text = par.runs[i + 1].text[rrlen + 1:]
                        return run
                        # print(run.text)
                elif len(par.runs[i + 1].text) + len(run.text) < ltag:
                    # error("wrong bookmark - " + tag + " in file - "+ template)
                    #############################
                    # hack - move text of run to next run
                    # ---------------------------
                    par.runs[i + 1].text = run.text + par.runs[i + 1].text
                    run.text = ""
        return None

    @staticmethod
    def write_table(document: Document, model, start_row=1):
        """fill table from model

        Args:
            doc, model

        Returns:

        """
        found = False
        for table in document.tables:
            #############################
            # search only first found
            # ---------------------------
            if found:
                continue
            if table:
                found = True
                # table.style = "table_text"
                #############################
                # fill table
                # ---------------------------
                # for i, r in enumerate(table.rows, sqlstart=-start_row):
                #     if i >= start_row:
                #         debug(r.cells[0].text)
                #         for k, cell in enumerate(r.cells):
                #             try:
                #                 cell.text= str(model[i][k])
                #             except IndexError:
                #                 error("out of range")
                # styles = document.styles
                # style = styles.add_style('table_nums', WD_STYLE_TYPE.PARAGRAPH)
                # style.base_style = styles['Normal']
                for i, mr in enumerate(model):
                    r = table.rows[i + start_row]
                    for k, cell in enumerate(r.cells):
                        try:
                            # cell.clear()
                            # cell.add_paragraph(str(mr[k]), "table_text")
                            cell.text = str(mr[k])
                            cell.paragraphs[0].style = "table_text"
                            if k >= 2:  # from 0
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                # cell.add_paragraph(str(mr[k]),style)
                            # else:
                        except IndexError:
                            error("out of range")
                    table.add_row()
                #############################
                # delete last row
                # ---------------------------
                table._tbl.remove(table.rows[-1]._tr)
                #############################
                # add borders
                # ---------------------------
                for cell in table._tbl.iter_tcs():
                    tcPr = cell.tcPr  # get tcPr element, in which we can define style of borders
                    tcBorders = OxmlElement('w:tcBorders')
                    top = OxmlElement('w:top')
                    top.set(qn('w:val'), 'single')

                    left = OxmlElement('w:left')
                    left.set(qn('w:val'), 'single')

                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '1')
                    bottom.set(qn('w:space'), '0')
                    bottom.set(qn('w:color'), '000000')

                    right = OxmlElement('w:right')
                    right.set(qn('w:val'), 'single')

                    tcBorders.append(top)
                    tcBorders.append(left)
                    tcBorders.append(bottom)
                    tcBorders.append(right)
                    tcPr.append(tcBorders)
                # table.style = "TableGrid"

    @staticmethod
    def fill_docx(taglist, table, template=None, out='./templates/results/new-file-name.docx'):
        document = Document(template)
        if table:
            PrepareDocument.write_table(document, table, 2)
        #############################
        # replace txt to totxt and save
        # ---------------------------
        PrepareDocument.fill_tags(taglist, document, True)
        out = out.replace("/", "_")
        document.save(out)
        info("file saved to - %s", out)
        #############################
        # open file
        # ---------------------------
        if hasattr(os, "startfile"):
            # windows
            os.startfile(out)
        else:
            # linux
            subprocess.call(['xdg-open', out])
        # Mac OS X ?

    @staticmethod
    def fill_max_to_pay(contr_set, end):
        sqlend = QDate(end).toString(SQL_DATE_FORMAT)
        contr_str = ", ".join(contr_set)
        # contr_str = contr_str[:-2]
        #############################
        # get max pay per month 50
        # ---------------------------
        max_to_pay_50 = PrepareDocument.get_max_to_pay(contr_str, sqlend, "50")
        #############################
        # get max pay per month 75
        # ---------------------------
        max_to_pay_75 = PrepareDocument.get_max_to_pay(contr_str, sqlend, "75")
        # pay_per_contr = []
        # for i, contr in enumerate(contr_set):
        #     pay_per_contr[i] =
        return None, max_to_pay_50, max_to_pay_50, max_to_pay_50, max_to_pay_75

    @staticmethod
    def get_max_to_pay(contr_str, sqlend, maxtp: str = "50"):
        #############################
        # prepare model
        # ---------------------------
        mpay = WD.models("max_pay_in_month_{}__where_for_fill".format(maxtp), "max_pay_in_month_{}".format(maxtp),
                         "max_pay_in_month_{2}.contracts_id in ({0}) and pddate <= '{1}'  order by pddate desc "
                         "".format(contr_str, sqlend, maxtp),
                         False)
        #############################
        # get data
        # ---------------------------
        if mpay.rowCount0() < 1:
            # app = QCoreApplication.instance()
            msg = "no data max_pay_in_month{}__where_for_fill".format(maxtp)
            critical(msg)
            # if isinstance(app, QApplication):
            #     # TODO: check pc avialable
            #     QMessageBox.critical(QApplication.topLevelWidgets()[0], "Error",
            #                          msg)
            return Decimal(0)
        max_to_pay = mpay.data(mpay.index(0, mpay.index_of_col("max_pay")), Qt.EditRole)
        #############################
        # check return
        # ---------------------------
        if max_to_pay is None:
            max_to_pay = 0
        max_to_pay = Decimal(max_to_pay)
        return max_to_pay

    @staticmethod
    def prepare_serv_table(client, start, end):
        sqlstart = QDate(start).toString(SQL_DATE_FORMAT)
        sqlend = QDate(end).toString(SQL_DATE_FORMAT)
        mserv = WD.models("_main_cprice__where_for_fill",
                          "_main_cprice",
                          "_main_cprice.client_id = {} and vdate between '{}' and '{}' ".format(
                              client, sqlstart, sqlend),  # dep,  and main_cprice.dep_id = {1}
                          False)
        header = ["contracts_id", "dep_id", "tnum", "serv", "serv_text", "quantity", "price", "to_pay", "servform_id",
                  "serv_id",
                  "sub_serv", "vdate"]  #
        table = mserv.read_into_2darray(header)
        # to_pay = Decimal(0)
        header[1] = "full_price"
        for row in table:
            row[header.index("full_price")] = row[header.index("quantity")] * row[header.index("price")]
            # to_pay = to_pay + Decimal(header.index("full_price"))
        table.header[1] = "full_price"
        return table

    @staticmethod
    def prepare_predv_serv_table(contr, dogv_start_date: QDate):
        # sqlstart = QDate(start).toString(SQL_DATE_FORMAT)
        # sqlend = QDate(end).toString(SQL_DATE_FORMAT)
        year = dogv_start_date.toString("yyyy")
        mserv = WD.models("_contr_has_serv_predv__where_for_fill",
                          "_contr_has_serv_predv",
                          "_contr_has_serv_predv.contracts_id = {} and `year` = {} ".format(contr, year),
                          False)
        header = ["contracts_id", "full_price", "tnum", "serv", "serv_text", "quantity", "price", "to_pay",
                  "servform_id",
                  "serv_id",
                  "sub_serv", "vdate"]  #
        table = mserv.read_into_2darray(header)
        # header = ["contracts_id", "dep_id", "tnum",
        #           "serv_text", "quantity", "price", "to_pay", "servform_id", "serv_id",
        #           "sub_serv", "vdate"]  #
        # to_pay = Decimal(0)
        # header[1] = "full_price"
        # for row in table:
        #     row[header.index("full_price")] = row[header.index("quantity")] * row[header.index("price")]
        #     # to_pay = to_pay + Decimal(header.index("full_price"))
        # table.header[1] = "full_price"
        return table

    @staticmethod
    def prepare_contr_model(contr_set):
        contr_mdl = tsQsfpModel(None, "_contracts")
        contr_mdl.super_model().fetch_all()
        contr_mdl.setFilterKeyColumn(contr_mdl.index_of_col("id"))
        str_dat = ["^" + str(x) + "$" for x in contr_set]
        re_pattern = "|".join(str_dat)
        # re: QRegularExpression = QRegularExpression("(" + re_pattern + ")")
        contr_mdl.setFilterRegularExpression(re_pattern)
        contr_mdl.sort(contr_mdl.index_of_col("enddate"))
        return contr_mdl

    @staticmethod
    @try_wrapper
    def print_doc(client, month, year, dep, contr="__all__", do_document="акт", vdate=None):
        start = 0
        end = 0
        if month == 13:
            start = dt(year, 1, 1)
            end = dt(year, 12, 31)
        else:
            start = dt(year, month, 1)
            fday, endday = calendar.monthrange(year, month)
            end = dt(year, month, endday)
        contr_count = 1
        contr_set = set()
        table = None
        cur_table = None
        remained_money = [-1, -1, -1, -1, -1]
        if do_document == "акт" or do_document == "квитанция":
            #############################
            # create serv table
            # ---------------------------
            table = PrepareDocument.prepare_serv_table(client, start, end)
            if not table:
                error("Нет услуг в этот период")
                return False
            #############################
            # fill_max_to_pay
            # ---------------------------
            contr_set = {str(x[0]) for x in table}
            remained_money = list(PrepareDocument.fill_max_to_pay(contr_set, end))
            #############################
            # prepare filtered contr model
            # ---------------------------
            if contr != "__all__":
                contr_set = set()
                contr_set.add(contr)
                # TODO: if several contrats in period - count all money
            contr_mdl = PrepareDocument.prepare_contr_model(contr_set)
            contr_count = contr_mdl.rowCount()
        elif do_document in ("предварительный расчет оплаты", "договор"):
            contr_set.add(contr)
            contr_mdl = PrepareDocument.prepare_contr_model(contr_set)
            contr_count = contr_mdl.rowCount()
            #############################
            # create serv table
            # ---------------------------
            contr_start = contr_mdl.data(contr_mdl.index(0, contr_mdl.index_of_col("startdate")), Qt.EditRole)
            table = PrepareDocument.prepare_predv_serv_table(contr, contr_start)
            #############################
            # fill_max_to_pay
            # ---------------------------
            contr_set = {str(x[0]) for x in table}
            remained_money = list(PrepareDocument.fill_max_to_pay(contr_set, end))
        elif contr and contr != "__all__":
            contr_set.add(contr)
            contr_mdl = PrepareDocument.prepare_contr_model(contr_set)
            contr_count = contr_mdl.rowCount()
        elif vdate:
            from data_worker import get_contract
            contr = get_contract(client, QDate(vdate).toString(SQL_DATE_FORMAT), dep)
            if contr == -1:
                debug("No contract for %s, %s, %s", client, QDate(vdate).toString(SQL_DATE_FORMAT), dep)
                pass  # TODO: other department
            contr_set.add(contr)
            contr_mdl = PrepareDocument.prepare_contr_model(contr_set)
            contr_count = contr_mdl.rowCount()
        #############################
        # call PrepareDocument.print_act_period for every contract
        # ---------------------------
        tempdir = tempfile.mkdtemp(None, "3uson")
        # TODO: also new act if percents changed
        for i in range(contr_count):
            cend = end
            cstart = start
            cur_contr = contr
            if do_document == "акт" or do_document == "квитанция":
                cur_contr = contr_mdl.data(contr_mdl.index(i, contr_mdl.index_of_col("id")), Qt.EditRole)
            # if cur_contr == contr or contr == "__all__":
            #############################
            # var setup
            # ---------------------------
            dl = {}
            #############################
            # fill contract data
            # ---------------------------
            PrepareDocument.fill_contract_data(contr_mdl, dl, i)
            #############################
            # recalculate start and end of period with respect to contract
            # ---------------------------
            contr_end__ = contr_mdl.data(contr_mdl.index(i, contr_mdl.index_of_col("enddate")), Qt.EditRole)
            contr_start = contr_mdl.data(contr_mdl.index(i, contr_mdl.index_of_col("startdate")), Qt.EditRole)
            if cend > contr_end__:
                cend = contr_end__
            if cstart < contr_start:
                cstart = contr_start
            PrepareDocument.fill_act_dates(dl, start, cend)
            #############################
            # get add_info datamodel
            # ---------------------------
            PrepareDocument.fill_add_info(dl, client, cur_contr, cend)
            # TODO: warning if services not in any contract
            #############################
            # livemin
            # ---------------------------
            if do_document in ("акт", "квитанция"):
                sql = " SELECT live_min_p FROM live_min where lmdate <= '{}' order by lmdate desc ".format(
                    QDate(cend).toString(SQL_DATE_FORMAT))
                if dl["work_livemin"]:
                    sql = " SELECT live_min_w FROM live_min where lmdate <= '{}' order by lmdate desc ".format(
                        QDate(cend).toString(SQL_DATE_FORMAT))
                dlconv(dl, "prmin", SD.line_query(sql))
            if do_document in ("договор", "предварительный расчет оплаты"):
                sql = " SELECT live_min_p FROM live_min where lmdate <= '{}' order by lmdate desc ".format(
                    contr_start.toString(SQL_DATE_FORMAT))
                if dl["work_livemin"]:
                    sql = " SELECT live_min_w FROM live_min where lmdate <= '{}' order by lmdate desc ".format(
                        contr_start.toString(SQL_DATE_FORMAT))
                prmin = SD.line_query(sql)
                if not prmin:
                    QMessageBox.critical(UI.main_window, "Ошибка",
                                         "Не удалось найти прожиточный минимум за период <= {}".
                                         format(contr_start.toString("dd.MM.yyyy")))
                dlconv(dl, "prmin", prmin)
            if do_document in ("акт", "квитанция", "договор", "предварительный расчет оплаты"):
                #############################
                # max to pay
                # ---------------------------
                if remained_money[1] > 0:
                    dlconv(dl, "payment_max", remained_money[1])
                    # dl["payment_max"] = "{0:0.2f}".format(remained_money[1]).replace(".", ",")
                else:
                    # fallback
                    dlconv(dl, "payment_max", (dl["___sdd"] - (1.5 * dl["___prmin"])) / 2)
                if remained_money[4] > 0:
                    dlconv(dl, "payment_max_st", remained_money[4])
                else:
                    # fallback
                    dlconv(dl, "payment_max_st", dl["___sdd"] * 0.75)
                #############################
                # prepare subtable
                # ---------------------------
                # table.remove_duplicate_with_sum(["quantity", "full_price", "to_pay"],
                #                                 ["contracts_id", "servform_id", "sub_serv", "vdate"])
                header = table.header
                if do_document in ("акт", "квитанция"):
                    table = [x for x in table if
                             cend >= x[table.header.index("vdate")] >= cstart]
                table2 = TableByRowsWHeader(table)
                table2.header = header
                if do_document in ("акт", "квитанция"):
                    cur_table = PrepareDocument.order_table(table2, cur_contr)
                elif do_document in ("договор", "предварительный расчет оплаты"):
                    cur_table = PrepareDocument.order_table2(table2)
                    tot = float(cur_table[-1][-1].replace(",", "."))
                    dlconv(dl, "tot1", tot)
                    dlconv(dl, "payment_full_num2", tot * dl["___proc"])
                #############################
                # servform
                # ---------------------------
                servform___ = contr_mdl.data(contr_mdl.index(i, contr_mdl.index_of_col("servform_id")), Qt.EditRole)
                #############################
                # fill payment
                # ---------------------------
                money = PrepareDocument.fill_payment(dl, cur_table, remained_money[servform___])
                #############################
                # update remained_money
                # ---------------------------
                remained_money[servform___] = remained_money[servform___] - money
                if remained_money[servform___] < 0:
                    remained_money[servform___] = 0
            #############################
            # fill document
            # ---------------------------
            if do_document == "квитанция":
                cur_table = None
            dl["otd"] = WD.models("dep").data_by_id(dep, 1)
            dl["otd2"] = dl["otd"]
            dl["otd3"] = dl["otd"]
            dl["otd4"] = dl["otd"]
            # file = "./templates/квитанция.docx"
            dl["pay_num"] = dl["payment_num"]
            dl["pay_num2"] = dl["pay_num"]
            if do_document == "акт":
                pass
            elif do_document == "договор":
                dl["topay_predv"] = dl["payment_num"]
                dl["topay_predv_long"], dl["topay_predv_cent"] = PrepareDocument.num_to_string(dl["topay_predv"], 3)
                cur_table = None
                if servform___ == 1:
                    do_document = do_document + " на дому"
                    dl["svaddress"] = dl["addr"]
                elif servform___ == 2:
                    do_document = do_document + " полустационар"
            elif do_document == "предварительный расчет оплаты":
                dl["topay_predv"] = dl["payment_num"]
                dl["topay_predv_long"], dl["topay_predv_cent"] = PrepareDocument.num_to_string(dl["topay_predv"], 3)
            #############################
            # convert dl to taglist
            # ---------------------------
            taglist = {}
            for k, val in dl.items():
                key = PrepareDocument.dcr(k)
                taglist[key] = val
                debug(str(key) + str(val))
            debug(taglist)
            #############################
            # output file
            # ---------------------------
            out = os.path.join(
                tempdir, do_document + "_" + dl["fio2"] + "_" + str(dl["ndog"]) + ".docx"
            ).replace(" ", "_").replace("/", "_")
            #############################
            # create file
            # ---------------------------
            file = os.path.join(PROJECT_DIR, "templates", do_document + ".docx")
            PrepareDocument.fill_docx(taglist, cur_table, file, out)
        #############################
        # TODO: check if contract ended in period b - print 2 acts
        # ---------------------------

    @staticmethod
    def order_table(table, cur_contr, count_total_price=False):
        """

        header = ["contracts_id", "full_price", "tnum", "serv", "serv_text", "quantity", "price", "to_pay", "servform_id",
         "serv_id", "sub_serv"]  #

        :param count_total_price:
        :param table:
        :param cur_contr:
        :return:
        """

        #############################
        # prepare 1st
        # ---------------------------
        # coltst = header.index("full_price")
        # coltend = header.index("serv_id") + 1
        # cur_table = TableByRowsWHeader([l[coltst:coltend] for l in table if l[0] == cur_contr])
        header = table.header
        cur_table = TableByRowsWHeader(table)
        cur_table.header = header
        cur_table.sort(
            key=itemgetter(
                header.index("tnum")), reverse=True)
        cur_table.remove_duplicate_with_sum(["quantity", "full_price", "to_pay"],
                                            ["contracts_id", "servform_id", "sub_serv", "vdate"])
        #############################
        # prepare total table
        # ---------------------------
        mdl_tot = WD.models("_serv_total")
        tot_header = "id, tnum, serv, serv_text, year, sub_serv, sub_serv_str, price, " \
                     "archive, total, acronym, workload, content".replace(" ", "").split(",")
        tot_table = TableByRowsWHeader(mdl_tot.read_into_2darray(tot_header))
        tot_header[0] = "serv_id"
        tot_table.header = tot_header
        res_table = TableByRowsWHeader()
        res_table.header = header
        #############################
        # insert all total rows
        # ---------------------------
        for i, row in enumerate(tot_table):
            with tot_table(i) as ttrow:
                if ttrow.serv_id not in res_table.serv_id:
                    res_table.append([])
                    with res_table(-1) as r:
                        r.tnum = ttrow.tnum
                        r.serv = ttrow.serv
                        r.serv_text = ttrow.serv_text
                        r.serv_id = ttrow.serv_id
                        if ttrow.year != SD.last_year:
                            r.serv = ttrow.serv + " " + str(ttrow.year)
                        r.quantity = "X"
                        r.price = "X"
                        r.to_pay = "X"
                        r.full_price = "X"
                        if count_total_price:
                            pass  # TODO:
        #############################
        # insert services rows
        # ---------------------------
        used_serv_id = set()
        for i, row1 in enumerate(cur_table):
            with cur_table(i) as row:
                sup_serv_ind = res_table.row_by_col_value("serv_id", row.sub_serv)
                if sup_serv_ind and sup_serv_ind >= 0:
                    res_table.insert(sup_serv_ind + 1, [])
                    with res_table(sup_serv_ind + 1) as r:
                        r.tnum = row.tnum
                        r.serv = row.serv
                        r.serv_text = row.serv_text
                        # price, to_pay, full_price is always right
                        # if ttrow.year != SD.last_year:
                        #     r.serv = ttrow.serv + " " + str(ttrow.year)
                        r.quantity = row.quantity
                        r.price = row.price
                        r.to_pay = row.to_pay
                        r.full_price = row.full_price
                        r.serv_id = row.serv_id
                        # r.sub_serv = row.sub_serv
                        sup_ind = row.sub_serv
                        while sup_ind:
                            used_serv_id.add(row.serv_id)
                            sup_ind = tot_table.sub_serv[
                                tot_table.row_by_col_value("serv_id", sup_ind)]
                else:
                    pass
                used_serv_id.add(row.sub_serv)
        debug("done order_table 1")
        return PrepareDocument.order_table2(res_table, used_serv_id)

    @staticmethod
    def order_table2(res_table, used_serv_id=None):
        #############################
        # remove unused rows and cols
        # ---------------------------
        res_table1 = []
        res_table_only_num = []
        res_table_only_num_header = ["tnum", "serv_text", "quantity", "price", "full_price"]
        for x in res_table:
            if used_serv_id is None or x[res_table.header.index("serv_id")] in used_serv_id:
                col3 = x[res_table.header.index("quantity")]
                col4 = x[res_table.header.index("price")]
                col5 = x[res_table.header.index("full_price")]
                try:
                    col3 = int(col3)
                except ValueError:
                    pass
                if not isinstance(col4, str):
                    res_table_only_num.append([0, 0, col3, col4, col5])
                    col4 = "{0:0.2f}".format(x[res_table.header.index("price")]).replace(".", ",")
                    col5 = "{0:0.2f}".format(x[res_table.header.index("full_price")]).replace(".", ",")
                col1 = x[res_table.header.index("tnum")]
                col2 = x[res_table.header.index("serv_text")].replace("___", ", ")
                res_table1.append([
                    col1,
                    col2,
                    col3,
                    col4,
                    col5]
                )
        # res_table1 = [
        #     [x[res_table.header.index("tnum")],
        #      x[res_table.header.index("serv_text")].replace("___", ", "),
        #      int(x[res_table.header.index("quantity")]),
        #      "{0:0.2f}".format(x[res_table.header.index("price")]).replace(".", ","),
        #      "{0:0.2f}".format(x[res_table.header.index("full_price")]).replace(".", ",")
        #
        #      ] for x in res_table if used_serv_id is None or x[res_table.header.index("serv_id")] in used_serv_id]
        #############################
        # add total row
        # ---------------------------
        quantity_col = res_table_only_num_header.index("quantity")
        full_price_col = res_table_only_num_header.index("full_price")
        res_table1.append(
            [
                "",
                "Итого",
                sum([int(x[quantity_col]) for x in res_table_only_num]),
                "X",
                "{0:0.2f}".format(
                    sum([x[full_price_col] for x in res_table_only_num])
                ).replace(".", ","),
            ]
        )
        res_table2 = TableByRowsWHeader(res_table1)
        res_table2.header = "tnum,serv_text,quantity,price,full_price".split(",")
        #############################
        # round numbers
        # ---------------------------
        for el in res_table2:
            el[-1] = PrepareDocument.cut_num(el[-1])
            el[-2] = PrepareDocument.cut_num(el[-2])
        # for l in table:
        #     l[-1] = "{0:0.2f}".format(l[-1])
        return res_table2

    @staticmethod
    def fill_contract_data(contr_mdl, dl_, i=0):
        """fill contract data"""
        cur_contr = contr_mdl.data(contr_mdl.index(i, contr_mdl.index_of_col("id")), Qt.EditRole)
        contr_end__ = contr_mdl.data(contr_mdl.index(i, contr_mdl.index_of_col("enddate")), Qt.EditRole)
        contr_start: QDate = contr_mdl.data(contr_mdl.index(i, contr_mdl.index_of_col("startdate")), Qt.EditRole)
        dl_["ndog"] = contr_mdl.data(contr_mdl.index(i, contr_mdl.index_of_col("contracts")), Qt.EditRole)
        dl_["contract_date"] = QDate(contr_start).toString("dd.MM.yyyy")
        dl_["dogvdate"] = dl_["contract_date"]
        dl_["day"] = contr_start.toString("dd")
        dl_["month"] = contr_start.toString("MMMM").lower(
        ).replace("ь", "я").replace("рт", "рта").replace("май", "мая").replace("август", "августа")
        dl_["year"] = contr_start.toString("yyyy")
        dl_["day"] = contr_start.toString("dd")
        dl_["dog_month"] = contr_start.toString("MMMM").lower(
        ).replace("ь", "я").replace("рт", "рта").replace("май", "мая").replace(
            "август", "августа")
        dl_["dog_year"] = contr_start.toString("yyyy")
        dl_["month_end"] = contr_end__.toString("MMMM").lower(
        ).replace("ь", "я").replace("рт", "рта").replace("май", "мая").replace(
            "август", "августа")
        dl_["year_end"] = contr_end__.toString("yyyy")
        dl_["day_end"] = contr_end__.toString("dd")
        dl_["dogv_stop"] = contr_end__.toString("dd.MM.yyyy")
        if contr_mdl.super_model()._relations_used:
            dl_["ripso"] = contr_mdl.data(contr_mdl.index(i, contr_mdl.index_of_col("ripso_id")), Qt.DisplayRole)
        else:
            dl_["ripso"] = contr_mdl.data(contr_mdl.index(i, contr_mdl.index_of_col("ripso_id")), Qt.EditRole)
            dl_["ripso"] = WD.get_data_from_model_name("ripso", "ripso",
                                                       dl_["ripso"])  # contr_mdl not a relational model
        if contr_end__ is None:
            warning("contract {} has no end date!".format(cur_contr))
        # elif end > contr_end__:
        #     end = contr_end__
        if contr_start is None:
            critical("contract {} has no start date!".format(cur_contr))
        # elif start < contr_start:
        #     start = contr_start

    @staticmethod
    def fill_add_info(dl, client, contr, end):
        """get add_info datamodel"""
        sqlend = end
        if not isinstance(end, str):
            sqlend = QDate(end).toString(SQL_DATE_FORMAT)
        #############################
        # get fio
        # ---------------------------
        clientmdl = WD.models("client")
        dl["telephone"] = clientmdl.data_by_id(client, clientmdl.index_of_col("phone"))
        dl["birthdate"] = clientmdl.data_by_id(client, clientmdl.index_of_col("clientbirth")).toString("dd.MM.yyyy")
        dl["snils"] = clientmdl.data_by_id(client, clientmdl.index_of_col("snils"))

        #############################
        # model add_info
        # ---------------------------
        _where = "add_info.contracts_id = {0} and pddate <= '{1}'".format(contr, sqlend)
        # mdl = WD.models("add_info__where_for_fill", "add_info", _where, False)
        mdl = tsQsfpModel(parent=WD, model_name="add_info__where_for_fill", tname="add_info", where=_where, rel=False)
        mdl.sort(mdl.super_model().fieldIndex("pddate"), Qt.DescendingOrder)
        ai_id = 1  # it sorted in descending order, but 0 is new row
        if mdl.rowCount0() > 1:
            # critical("Сan't get metadata for print!!!")
            fio_from_ai = mdl.data_rc(ai_id, mdl.index_of_col("curFIO"))
            if len(fio_from_ai) > 5:  # sanity check
                dl["fio"] = fio_from_ai
        #############################
        # fio 2
        # ---------------------------
        dl["fio"] = clientmdl.data_by_id(client, clientmdl.index_of_col("client"), 0, "client_id")
        dl["fio2"] = dl["fio"]
        dl["fiocrop"] = PrepareDocument.cropFIO(dl["fio"])
        dl["fiocrop2"] = dl["fiocrop"]
        #############################
        # exit if no add_info
        # ---------------------------
        if mdl.rowCount0() < 1:
            critical("Сan't get metadata for print!!!")
            return False
        dl["work_livemin"] = mdl.data_rc(ai_id, mdl.index_of_col("work_livemin"))
        dl["addr"] = mdl.data_rc(ai_id, mdl.index_of_col("address"))
        dl["pasp"] = mdl.data_rc(ai_id, mdl.index_of_col("psp"))
        dl["proc"] = mdl.data_rc(ai_id, mdl.index_of_col("perc"))
        dlconv(dl, "proc", dl["proc"])
        dl["proc2"] = dl["proc"]
        aver_income = mdl.data_rc(ai_id, mdl.index_of_col("sdd"))
        dlconv(dl, "sdd", aver_income)
        dl["passport"] = dl["pasp"]
        if dl["addr"] == "":
            dl["addr"] = "____________________________________"
        dl["addr2"] = dl["addr"]
        dl["address"] = dl["addr"]
        dl["address2"] = dl["addr"]
        if dl["passport"] == "":
            dl["passport"] = "____________________________________________"
        dl["passport2"] = dl["passport"]
        dl["pasport"] = dl["passport"]
        dl["pasport2"] = dl["passport"]
        dl["aver_income"] = PrepareDocument.num_to_string(aver_income)
        dl["aver_income_num"] = "{0:0.2f}".format(aver_income).replace(".", ",")
        # return mdl

    @staticmethod
    def fill_act_dates(dl, start, end):
        #############################
        # dates
        # ---------------------------
        dl["from_date"] = QDate(start).toString("dd.MM.yyyy")
        dl["to_date"] = QDate(end).toString("dd.MM.yyyy")
        dl["act_date"] = (QDate(end).addDays(1)).toString("dd.MM.yyyy")

    @staticmethod
    def fill_payment(dl: dict, table, max_to_pay=None):
        #############################
        # get services per contract
        # ---------------------------
        to_pay = 0
        full_price = 0
        for i, r in enumerate(table):
            with table(i) as row:
                try:
                    serv_tot_price = row.quantity * row.price
                    full_price += serv_tot_price
                    to_pay += (serv_tot_price * float(dl["___proc"]))
                except TypeError:
                    try:
                        serv_tot_price = row.quantity * float(row.price.replace(",", "."))
                        full_price += serv_tot_price
                        to_pay += (serv_tot_price * float(dl["___proc"]))
                    except (ValueError, TypeError):
                        pass
        #############################
        # check max per month
        # ---------------------------
        to_pay = Decimal(to_pay)
        if max_to_pay >= 0:
            if max_to_pay < to_pay:
                to_pay = max_to_pay
        #############################
        # fill data dl
        # ---------------------------
        dl["payment_full_num"] = "{0:0.2f}".format(full_price).replace(".", ",")
        dl["payment_full"] = PrepareDocument.num_to_string(full_price)
        dl["payment_num"] = "{0:0.2f}".format(float(to_pay)).replace(".", ",")
        dl["payment"] = PrepareDocument.num_to_string(float(to_pay))
        dl["proc"] = str(int(dl["___proc"] * 100))  # + "%"
        dl["proc2"] = dl["proc"]
        return to_pay


def main():
    try:
        os.mkdir(os.path.join(PROJECT_DIR, "templates", "results"))
    except FileExistsError:
        pass
    print("test file - ./templates/Акт предоставления услуг общий.docx")
    taglist = {'$__ndog__': '338/2019', '$__contract_date__': '01.09.2019', '$__dogvdate__': '01.09.2019',
               '$__day__': '01',
               '$__month__': 'Сентябрь', '$__year__': '2019', '$__dogv_stop__': '01.09.2022',
               '$__from_date__': '01.09.2019',
               '$__to_date__': '30.09.2019', '$__act_date__': '01.10.2019', '$__telephone__': '',
               '$__birthdate__': '28.11.1935',
               '$__snils__': '006-580-139 29', '$__fio__': 'Андреянова Валентина Ивановна',
               '$__fio2__': 'Андреянова Валентина Ивановна', '$__fiocrop__': 'Андреянова В.И.',
               '$__fiocrop2__': 'Андреянова В.И.', '$__addr__': '4-я Красноармейская ул.,д.18Б,кв.12',
               '$__addr2__': '4-я Красноармейская ул.,д.18Б,кв.12',
               '$__address__': '4-я Красноармейская ул.,д.18Б,кв.12',
               '$__address2__': '4-я Красноармейская ул.,д.18Б,кв.12',
               '$__passport__': 'паспорт 4000 546548 выдан 38 ОМ Адмиралтейского района 15.06.2001 ',
               '$__passport2__': 'паспорт 4000 546548 выдан 38 ОМ Адмиралтейского района 15.06.2001 ',
               '$__aver_income__': 'двадцать четыре тысячи семьсот тридцать два руб. 26 коп.',
               '$__aver_income_num__': '24732,26',
               '$__proc__': '20', '$__payment_full_num__': '1392,00',
               '$__payment_full__': 'одна тысяча триста девяносто два руб. 0 коп.', '$__payment_num__': '278,40',
               '$__payment__': 'двести семьдесят восемь руб. 40 коп.'}

    # table = [[1, 2, 3], [4, 5, 6], [7, 8, 9]]
    print(taglist)
    PrepareDocument.console_connect()
    # PrepareDocument.fill_docx(taglist, table, "./templates/акт.docx")
    #############################
    # run
    # ---------------------------
    # PrepareDocument.print_doc(1, 6, 2001, 1)
    # PrepareDocument.print_doc(69, 1, 2019, 31, contr="__all__", do_document="акт", vdate=None)
    PrepareDocument.print_doc(69, 13, 2019, 31, "1485", do_document="предварительный расчет оплаты", vdate=None)
    debug("done")


if __name__ == "__main__":
    main()
